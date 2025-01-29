import streamlit as st
import pandas as pd
import camelot
import PyPDF2
import tempfile
import os
import base64
import numpy as np
from fuzzywuzzy import process

# Bibliotecas para gerar DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Biblioteca para gerar PDF
from fpdf import FPDF

import re


# ==========================================
#   CONFIGURAÇÃO INICIAL DO STREAMLIT
# ==========================================
st.set_page_config(page_title="Analista de Extratos Bancários", layout="centered")

logo_path = "MP.png"           # Caminho para a logomarca
glossary_path = "Tarifas.txt"  # Caminho para o glossário


# ==========================================
#   FUNÇÃO PARA SANITIZAR O NOME DO CLIENTE
# ==========================================
def sanitize_nome_cliente(nome: str) -> str:
    """
    Remove acentos/caracteres especiais e troca espaços por underscores,
    para evitar problemas em nomes de arquivos.
    """
    nome = nome.strip()
    nome = nome.replace(" ", "_")
    nome = re.sub(r"[^\w\-_]", "", nome)
    return nome


# ==========================================
#   FUNÇÃO PARA TENTAR EXTRAIR NOME DO CLIENTE
# ==========================================
def extrair_nome_cliente(pdf_path):
    """
    Tenta extrair o nome do cliente de dentro do PDF,
    procurando pela substring 'Nome:' e pegando o conteúdo até uma quebra de linha.
    Caso não encontre, retorna 'Sem_Nome'.
    """
    try:
        with open(pdf_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            first_page_text = pdf_reader.pages[0].extract_text()

            if "Nome:" in first_page_text:
                pos = first_page_text.find("Nome:") + len("Nome:")
                restante = first_page_text[pos:].strip()
                linha = restante.split("\n")[0].strip()
                return sanitize_nome_cliente(linha) if linha else "Sem_Nome"
    except:
        pass

    return "Sem_Nome"


# ==========================================
#   FUNÇÃO PARA FORMATAR VALORES EM R$
# ==========================================
def formatar_valor_brl(valor):
    """
    Converte um valor numérico (float/int) para o formato BRL,
    usando ponto para milhar e vírgula para decimais, ex: 1.234,56.
    Se não for conversível em float, retorna como string original.
    """
    try:
        val = float(valor)
        # Formatação no estilo 1.234,56
        # 1) formata no estilo "1,234.56"
        txt_temp = f"{val:,.2f}"
        # 2) troca as vírgulas por pontos e pontos por vírgulas
        txt_temp = txt_temp.replace(",", "X").replace(".", ",").replace("X", ".")
        return txt_temp
    except:
        return str(valor)


# ==========================================
#   FUNÇÃO PARA GERAR PDF EM PAISAGEM
# ==========================================
def df_to_pdf_bytes(df, titulo="Relatório", formatar_linhas_especiais=False):
    """
    Gera um PDF (em bytes) no formato A4 'landscape' (paisagem),
    com a coluna "Histórico" suficientemente larga para caber em uma única linha.
    Opcionalmente, formata linhas especiais como "Valor Total (R$)" e "Em dobro (R$)".
    """

    class PDFTabela(FPDF):
        def __init__(self, orientation='L', unit='mm', format='A4', formatar_linhas_especiais=False):
            super().__init__(orientation, unit, format)
            self.set_left_margin(10)
            self.set_right_margin(10)
            self.set_top_margin(10)
            self.set_auto_page_break(auto=True, margin=10)

            self.title_str = titulo

            # Colunas
            self.col_names = df.columns.tolist()
            # Larguras pré-definidas para cada coluna em modo paisagem (total: 277 mm)
            self.col_widths = []
            for col in self.col_names:
                if col == "Histórico":
                    self.col_widths.append(210)
                elif col in ["Débito (R$)", "Crédito (R$)"]:
                    self.col_widths.append(27)
                else:
                    self.col_widths.append(20)

            # Alinhamentos
            self.col_aligns = []
            for col in self.col_names:
                if col in ["Data", "Docto."]:
                    self.col_aligns.append("C")
                elif col == "Histórico":
                    self.col_aligns.append("L")
                else:
                    self.col_aligns.append("R")

            self.row_height = 7
            self.font_size_data = 9
            self.font_size_header = 10
            self.font_size_title = 14
            self.formatar_linhas_especiais = formatar_linhas_especiais

        def header(self):
            self.set_font("Arial", "B", self.font_size_title)
            self.cell(0, 10, self.title_str, 0, 1, "C")
            self.ln(2)

            self.set_font("Arial", "B", self.font_size_header)
            for i, col_name in enumerate(self.col_names):
                w = self.col_widths[i]
                a = self.col_aligns[i]
                self.cell(w, self.row_height, str(col_name), border=1, ln=0, align=a)
            self.ln(self.row_height)

        def footer(self):
            self.set_y(-15)
            self.set_font("Arial", "I", 8)
            self.cell(0, 10, f"Página {self.page_no()}/{{nb}}", 0, 0, "C")

        def gerar_tabela(self):
            self.set_font("Arial", "", self.font_size_data)
            for _, row in df.iterrows():
                if self.get_y() + self.row_height > (self.h - 10):
                    self.add_page()

                # Verifica se a linha é especial
                is_especial = (
                    "Histórico" in row and row["Histórico"] in ["Valor Total (R$)", "Em dobro (R$)"]
                )
                if is_especial and self.formatar_linhas_especiais:
                    self.set_font("Arial", "B", 14)
                    self.set_text_color(255, 0, 0)  # Vermelho
                else:
                    self.set_font("Arial", "", self.font_size_data)
                    self.set_text_color(0, 0, 0)

                for i, col_name in enumerate(self.col_names):
                    txt = str(row[col_name])
                    w = self.col_widths[i]
                    a = self.col_aligns[i]
                    self.cell(w, self.row_height, txt, border=1, ln=0, align=a)
                self.ln(self.row_height)

                if is_especial and self.formatar_linhas_especiais:
                    self.set_font("Arial", "", self.font_size_data)
                    self.set_text_color(0, 0, 0)

    if df.empty:
        pdf_vazio = FPDF(orientation='L', unit='mm', format='A4')
        pdf_vazio.add_page()
        pdf_vazio.set_font("Arial", "B", 12)
        pdf_vazio.cell(0, 10, "DataFrame vazio - nenhum dado para exibir.", 0, 1, "C")
        return pdf_vazio.output(dest="S").encode("latin-1")

    pdf = PDFTabela(formatar_linhas_especiais=formatar_linhas_especiais)
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.gerar_tabela()
    return pdf.output(dest="S").encode("latin-1")


# ==========================================
#   FUNÇÃO PARA GERAR DOCX EM PAISAGEM
# ==========================================
def df_to_doc_bytes(df, titulo="Relatório", adicionar_totais=False):
    """
    Gera um documento DOCX no formato paisagem,
    com a coluna "Histórico" suficientemente larga para caber em uma única linha.
    Se adicionar_totais=True, insere ao final as linhas de "Valor Total (R$)" e "Em dobro (R$)".
    """
    document = Document()

    # Configurar orientação da página para paisagem
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # Adicionar título
    title = document.add_heading(titulo, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Se DataFrame vazio, apenas DOCX básico
    if df.empty:
        p = document.add_paragraph("DataFrame vazio - nenhum dado para exibir.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buffer_empty = BytesIO()
        document.save(buffer_empty)
        return buffer_empty.getvalue()

    # Adicionar tabela
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Largura das colunas (heurística)
    col_widths_inches = []
    for col in df.columns:
        if col == "Histórico":
            col_widths_inches.append(227 / 25.4)
        elif col in ["Débito (R$)", "Crédito (R$)"]:
            col_widths_inches.append(20 / 25.4)
        else:
            col_widths_inches.append(15 / 25.4)
    for i, width in enumerate(col_widths_inches):
        table.columns[i].width = Inches(width)

    # Adicionar dados
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            if df.columns[i] == "Histórico":
                text = str(item).replace('\n', ' ').replace('\r', ' ').strip()
            else:
                text = str(item)

            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.add_run(text)

            if df.columns[i] == "Histórico":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run.font.size = Pt(9)
            paragraph.line_spacing = Pt(9)

            # Se for linha "Valor Total (R$)" ou "Em dobro (R$)", destacar
            if "Histórico" in row and row["Histórico"] in ["Valor Total (R$)", "Em dobro (R$)"]:
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 0, 0)

    # Se quisermos que a própria função gere as linhas de soma
    if adicionar_totais:
        # Verificar se existe "Débito (R$)" ou "Crédito (R$)" para somar
        if "Débito (R$)" in df.columns:
            total_col = "Débito (R$)"
        else:
            total_col = "Crédito (R$)"

        # Converter valores
        numeros = (
            df[total_col]
            .str.replace("R$", "", regex=False)   # remove 'R$'
            .str.replace(" ", "", regex=False)    # remove espaços
            .str.replace(".", "", regex=False)    # remove pontos
            .str.replace(",", ".", regex=False)   # troca vírgula por ponto
        )
        total = pd.to_numeric(numeros, errors='coerce').sum()

        # Linha de Valor Total
        total_row = table.add_row().cells
        if len(total_row) > 1:
            total_row[1].text = "Valor Total (R$)"
        if len(total_row) > 3:
            total_row[3].text = f"{total:.2f}"
        for i, cell in enumerate(total_row):
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            if i == 1 or i == 3:
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 0, 0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT

        # Linha de Em dobro (R$)
        double_total = total * 2
        double_row = table.add_row().cells
        if len(double_row) > 1:
            double_row[1].text = "Em dobro (R$)"
        if len(double_row) > 3:
            double_row[3].text = f"{double_total:.2f}"
        for i, cell in enumerate(double_row):
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            if i == 1 or i == 3:
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 0, 0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ==========================================
#   FUNÇÕES AUXILIARES
# ==========================================
def get_image_base64(file_path):
    """
    Converte uma imagem para base64 para exibição no Streamlit.
    """
    if not os.path.exists(file_path):
        st.warning(f"Logomarca não encontrada em: {file_path}")
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


def carregar_glossario(path):
    """
    Carrega o glossário de um arquivo texto, 1 termo por linha.
    """
    try:
        with open(path, "r", encoding="utf-8") as file:
            return file.read().splitlines()
    except IOError:
        st.error(f"Erro ao ler o arquivo de glossário em: {path}")
        return []


def match_glossary(text, glossary, threshold=85):
    """
    Verifica se um texto corresponde a algum termo no glossário
    com base no limiar de similaridade (0-100).
    """
    if not glossary or not text:
        return False
    result = process.extractOne(text, glossary)
    return (result is not None) and (result[1] >= threshold)


def filtrar_por_glossario(df, glossary, threshold=85):
    """
    Filtra o DataFrame pelo glossário (similaridade >= threshold).
    """
    if df.empty or not glossary:
        return pd.DataFrame()
    mask = df["Histórico"].apply(lambda x: match_glossary(x, glossary, threshold))
    return df[mask]


def obter_numero_de_paginas(pdf_path):
    with open(pdf_path, "rb") as file:
        pdf_reader = PyPDF2.PdfReader(file)
        return len(pdf_reader.pages)


def ignorar_tabela(df):
    """
    Verifica se deve ignorar a tabela (linhas de propaganda ou outro texto).
    """
    condicao = (
        (df == "Fone Fácil Bradesco").any(axis=1)
        | (df == "Se Preferir, fale com a BIA pelo").any(axis=1)
        | (df == "Saldo Invest Fácil").any(axis=1)
    )
    return not (condicao == False).all()


def converter_data_para_dois_digitos(data):
    """
    Converte o ano de 4 dígitos para 2 dígitos no formato DD/MM/AA.
    Exemplo: 01/01/2023 -> 01/01/23
    """
    if pd.isna(data) or data == "":
        return data
    try:
        partes = data.split('/')
        if len(partes) == 3 and len(partes[2]) == 4:
            partes[2] = partes[2][2:]
            return '/'.join(partes)
        return data
    except:
        return data


def processar_pdf(pdf_path):
    """
    Processa o PDF com o extrato bancário e gera um DataFrame consolidado.
    """
    try:
        num_paginas = obter_numero_de_paginas(pdf_path)
        formato = ["90, 220, 320, 420, 520"] * num_paginas

        tables = camelot.read_pdf(
            pdf_path,
            pages="all",
            row_tol=15,
            flavor="stream",
            columns=formato,
        )

        extrato = pd.DataFrame(columns=["Data", "Histórico", "Docto.", "Crédito (R$)", "Débito (R$)", "Saldo (R$)"])

        with st.spinner('Processando tabelas...'):
            progress_bar = st.progress(0)
            for i, table in enumerate(tables):
                df = table.df
                if ignorar_tabela(df):
                    continue

                # Identifica linha de cabeçalho
                check_start = (df == "Data").any(axis=1)
                if any(check_start):
                    idx = check_start.idxmax()
                    df = df[idx + 1:]

                # Ajustar colunas
                expected_cols = ["Data", "Histórico", "Docto.", "Crédito (R$)", "Débito (R$)", "Saldo (R$)"]
                if len(df.columns) >= 6:
                    df.columns = expected_cols
                else:
                    df = df.iloc[:, :6]
                    df.columns = expected_cols

                extrato = pd.concat([extrato, df], ignore_index=True)
                progress_bar.progress((i + 1) / len(tables))

        # Ajustar datas
        extrato["Data"] = extrato["Data"].replace("", np.nan).ffill()
        extrato["Data"] = extrato["Data"].apply(converter_data_para_dois_digitos)

        # Linhas em branco (juntar histórico)
        linhas_vazias = extrato[
            (extrato["Docto."] == "") &
            (extrato["Crédito (R$)"] == "") &
            (extrato["Débito (R$)"] == "") &
            (extrato["Saldo (R$)"] == "")
        ].index
        for idx in linhas_vazias:
            if idx + 1 < len(extrato):
                extrato.at[idx + 1, "Histórico"] = (
                    extrato.at[idx, "Histórico"] + " " + extrato.at[idx + 1, "Histórico"]
                )

        extrato = extrato.drop(linhas_vazias).reset_index(drop=True)
        return extrato

    except Exception as e:
        st.error(f"Erro ao processar o PDF: {str(e)}")
        return None


def filtrar_debitos(df):
    """
    Filtra somente as linhas que contêm Débito (R$) preenchido.
    Remove colunas de crédito/saldo.
    """
    debitos = df[df["Débito (R$)"].notna() & (df["Débito (R$)"] != "")]
    cols_drop = ["Crédito (R$)", "Saldo (R$)"]
    cols_presentes = [c for c in cols_drop if c in debitos.columns]
    debitos = debitos.drop(columns=cols_presentes)
    return debitos


def filtrar_creditos(df):
    """
    Filtra somente as linhas que contêm Crédito (R$) preenchido.
    Remove colunas de débito/saldo.
    """
    creditos = df[df["Crédito (R$)"].notna() & (df["Crédito (R$)"] != "")]
    cols_drop = ["Débito (R$)", "Saldo (R$)"]
    cols_presentes = [c for c in cols_drop if c in creditos.columns]
    creditos = creditos.drop(columns=cols_presentes)
    return creditos


# ==========================================
#     MAIN STREAMLIT
# ==========================================
def main():
    # Inicializar keys
    for key in [
        "df_extrato", "df_debito", "df_debito_gloss", "df_debito_gloss_filtrado",
        "df_credito", "df_credito_gloss", "df_credito_gloss_filtrado",
        "operacao_selecionada", "nome_cliente"
    ]:
        if key not in st.session_state:
            st.session_state[key] = None

    # Exibir logomarca
    image_base64 = get_image_base64(logo_path)
    if image_base64:
        st.markdown(
            f"""
            <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 20px;">
                <img src="data:image/png;base64,{image_base64}" alt="Logomarca" style="width: 300px;">
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.subheader("Análise de Extratos Bancários do Bradesco")

    # Carrega glossário
    glossary_terms = carregar_glossario(glossary_path)
    if not glossary_terms:
        st.warning("Glossário não encontrado ou vazio!")

    # Upload do PDF
    uploaded_file = st.file_uploader("Inserir Extrato Bancário do Bradesco", type="pdf")
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            pdf_path = tmp_file.name

        # Tentar extrair nome do cliente
        nome_cliente_encontrado = extrair_nome_cliente(pdf_path)
        st.session_state["nome_cliente"] = nome_cliente_encontrado

        try:
            df_extrato = processar_pdf(pdf_path)
            st.session_state["df_extrato"] = df_extrato

            if df_extrato is not None and not df_extrato.empty:
                st.success("PDF processado com sucesso!")
                st.markdown("### Extrato Completo")
                st.dataframe(df_extrato, use_container_width=True)
            else:
                st.warning("Não foi possível processar o PDF ou o extrato está vazio.")
        finally:
            os.unlink(pdf_path)

    operacao = st.session_state.get("operacao_selecionada", None)
    nome_cliente = st.session_state.get("nome_cliente", "Sem_Nome")

    if st.session_state.get("df_extrato") is not None and not st.session_state["df_extrato"].empty:
        # Se ainda não escolheu a operação (Débito/Crédito)
        if not operacao:
            st.markdown("### Qual operação deseja analisar?")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Débito"):
                    st.session_state["operacao_selecionada"] = "Débito"
                    operacao = "Débito"
            with col2:
                if st.button("Crédito"):
                    st.session_state["operacao_selecionada"] = "Crédito"
                    operacao = "Crédito"

        # ========================= DÉBITOS =========================
        if operacao == "Débito":
            st.markdown("## Análise de Débitos")
            with st.form("filtrar_debitos_form"):
                st.markdown("### 1) Filtrar Operações de Débito")
                filtrar_debitos_submit = st.form_submit_button("Filtrar Débitos")

            if filtrar_debitos_submit:
                df_debito = filtrar_debitos(st.session_state["df_extrato"])
                st.session_state["df_debito"] = df_debito
                st.markdown("#### Resultado (Extrato de Débito)")
                st.dataframe(df_debito, use_container_width=True)

                pdf_debitos = df_to_pdf_bytes(df_debito, titulo="Extrato de Débitos")
                st.download_button(
                    label="Baixar PDF (Débitos)",
                    data=pdf_debitos,
                    file_name=f"debitos_{nome_cliente}.pdf",
                    mime="application/pdf",
                )

            if st.session_state.get("df_debito") is not None and not st.session_state["df_debito"].empty:
                with st.form("filtrar_glossario_debito_form"):
                    st.markdown("### 2) Filtrar Débitos no Glossário (com Precisão Ajustável)")
                    precision_debito = st.slider(
                        "Precisão da correspondência para Débitos (0.5 a 1.0):",
                        min_value=0.5,
                        max_value=1.0,
                        value=0.85,
                        step=0.025
                    )
                    filtrar_gloss_debito_submit = st.form_submit_button("Filtrar Débitos no Glossário")

                if filtrar_gloss_debito_submit:
                    df_debito_gloss = filtrar_por_glossario(
                        st.session_state["df_debito"], glossary_terms, threshold=int(precision_debito * 100)
                    )
                    # Remove colunas extras se existirem
                    df_debito_gloss = df_debito_gloss.drop(columns=["Crédito (R$)", "Saldo (R$)"], errors='ignore')
                    st.session_state["df_debito_gloss"] = df_debito_gloss
                    st.session_state["df_debito_gloss_filtrado"] = None

                    st.markdown("#### Resultado: Débitos + Glossário")
                    st.dataframe(df_debito_gloss, use_container_width=True)

                    pdf_gloss_debito = df_to_pdf_bytes(df_debito_gloss, titulo="Débitos (Filtrados no Glossário)")
                    st.download_button(
                        label="Baixar PDF (Débitos Glossário)",
                        data=pdf_gloss_debito,
                        file_name=f"debitos_glossario_{nome_cliente}.pdf",
                        mime="application/pdf",
                    )

            if st.session_state.get("df_debito_gloss") is not None and not st.session_state["df_debito_gloss"].empty:
                with st.form("excluir_debitos_form"):
                    st.markdown("### 3) Lista Única de 'Histórico' para Débitos + Inclusão")

                    df_gloss_original_debito = st.session_state["df_debito_gloss"]
                    df_base_exclusao_debito = (
                        st.session_state["df_debito_gloss_filtrado"]
                        if st.session_state["df_debito_gloss_filtrado"] is not None
                        else df_gloss_original_debito
                    ).copy()

                    valores_unicos_debito = sorted(df_base_exclusao_debito["Histórico"].unique())
                    st.markdown("#### Lista Única de 'Histórico' (Débitos - sem repetições)")
                    st.write("Marque os itens que deseja incluir:")

                    selected_historicos_debito = []
                    for i, hist in enumerate(valores_unicos_debito):
                        count_hist = df_base_exclusao_debito[df_base_exclusao_debito["Histórico"] == hist].shape[0]
                        rotulo = f"{i+1}- {hist} ({count_hist} {'vez' if count_hist == 1 else 'vezes'})"
                        if st.checkbox(rotulo, key=f"unique_hist_debito_{i}"):
                            selected_historicos_debito.append(hist)

                    confirmar_inclusao_debito = st.form_submit_button("Confirmar Inclusão (Débitos)")
                    if confirmar_inclusao_debito:
                        if selected_historicos_debito:
                            df_filtrado_debito = df_base_exclusao_debito[
                                df_base_exclusao_debito["Histórico"].isin(selected_historicos_debito)
                            ]
                            df_filtrado_debito = df_filtrado_debito.reset_index(drop=True)

                            st.success("Operações de Débito incluídas com sucesso!")
                            st.session_state["df_debito_gloss_filtrado"] = df_filtrado_debito

                            st.markdown("#### Lista Restante após Inclusões (Débitos - sem repetições)")
                            if df_filtrado_debito.empty:
                                st.write("Nenhum histórico de Débito restante.")
                            else:
                                df_restante_unicos_debito = df_filtrado_debito["Histórico"].value_counts().reset_index()
                                df_restante_unicos_debito.columns = ["Histórico", "Ocorrências"]
                                st.dataframe(df_restante_unicos_debito, use_container_width=True)
                        else:
                            st.warning("Nenhuma descrição de Débito foi selecionada.")

            # --- APRESENTAÇÃO DOS DÉBITOS FINAIS (COM TOTAL POSITIVO) ---
            if st.session_state.get("df_debito_gloss_filtrado") is not None and not st.session_state["df_debito_gloss_filtrado"].empty:
                with st.form("apresentar_tarifas_debito_form"):
                    st.markdown("### 4) Apresentar Tarifas para Débitos (DataFrame Final Ordenado)")
                    apresentar_tarifas_debito_submit = st.form_submit_button("Apresentar Tarifas para Débitos")

                if apresentar_tarifas_debito_submit:
                    df_para_exibir_debito = st.session_state["df_debito_gloss_filtrado"]
                    if not df_para_exibir_debito.empty:
                        # Remover colunas não necessárias
                        df_para_exibir_debito = df_para_exibir_debito.drop(columns=["Crédito (R$)", "Saldo (R$)"], errors='ignore')

                        # Converter valores de Débito para float
                        numeros_debitos = (
                            df_para_exibir_debito["Débito (R$)"]
                            .str.replace("R$", "", regex=False)
                            .str.replace(" ", "", regex=False)
                            .str.replace(".", "", regex=False)
                            .str.replace(",", ".", regex=False)
                        )
                        valores_float_debito = pd.to_numeric(numeros_debitos, errors='coerce').fillna(0.0)

                        # ------------------------------------------------------------------
                        # AQUI FORÇAMOS QUE O SOMATÓRIO SEJA POSITIVO
                        # (Podemos simplesmente usar abs() em cada linha ou multiplicar por -1
                        #  se as linhas já estiverem negativas. Abaixo uso abs() para garantir.)
                        valores_float_debito = valores_float_debito.abs()
                        # ------------------------------------------------------------------

                        # Soma total (positivo)
                        total_debitos = valores_float_debito.sum()

                        # Cria colunas já formatadas em BRL
                        df_para_exibir_debito["Débito (R$)"] = valores_float_debito.apply(formatar_valor_brl)

                        # Adicionar duas linhas de total e em dobro (ambos positivos)
                        valor_total = pd.DataFrame({
                            "Data": [""],
                            "Histórico": ["Valor Total (R$)"],
                            "Docto.": [""],
                            # Multiplicar por -1 se quisesse "reverter", mas aqui já está tudo positivo
                            "Débito (R$)": [formatar_valor_brl(total_debitos)]
                        })
                        em_dobro = pd.DataFrame({
                            "Data": [""],
                            "Histórico": ["Em dobro (R$)"],
                            "Docto.": [""],
                            "Débito (R$)": [formatar_valor_brl(total_debitos * 2)]
                        })

                        extrato_debito_final = pd.concat([df_para_exibir_debito, valor_total, em_dobro], ignore_index=True)

                        # Ordenar por data (mantendo totais no fim)
                        extrato_debito_final["Data"] = pd.to_datetime(
                            extrato_debito_final["Data"], format="%d/%m/%y", errors='coerce'
                        )
                        extrato_debito_final = extrato_debito_final.sort_values(by="Data")
                        extrato_debito_final["Data"] = extrato_debito_final["Data"].dt.strftime("%d/%m/%y")
                        extrato_debito_final["Data"] = extrato_debito_final["Data"].fillna("")

                        # Linhas de total devem ter a data vazia
                        extrato_debito_final.loc[
                            extrato_debito_final["Histórico"].isin(["Valor Total (R$)", "Em dobro (R$)"]),
                            "Data"
                        ] = ""

                        st.markdown("#### DataFrame Final de Débitos (Cronológico)")
                        st.dataframe(extrato_debito_final, use_container_width=True)

                        # PDF com destaque nos totais
                        pdf_final_debito = df_to_pdf_bytes(
                            extrato_debito_final,
                            titulo="Extrato Final de Débitos (Cronológico)",
                            formatar_linhas_especiais=True
                        )
                        st.download_button(
                            label="Baixar PDF (Débitos Final - Cronológico)",
                            data=pdf_final_debito,
                            file_name=f"debitos_final_cronologico_{nome_cliente}.pdf",
                            mime="application/pdf",
                        )

                        # DOCX (já temos as linhas de total, então adicionar_totais=False)
                        doc_final_debito = df_to_doc_bytes(
                            extrato_debito_final,
                            titulo="Extrato Final de Débitos (Cronológico)",
                            adicionar_totais=False
                        )
                        st.download_button(
                            label="Baixar DOCX (Débitos Final - Cronológico)",
                            data=doc_final_debito,
                            file_name=f"debitos_final_cronologico_{nome_cliente}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    else:
                        st.warning("Não há extrato final para Débitos.")

        # ========================= CRÉDITOS =========================
        elif operacao == "Crédito":
            st.markdown("## Análise de Créditos")
            with st.form("filtrar_creditos_form"):
                st.markdown("### 1) Filtrar Operações de Crédito")
                filtrar_creditos_submit = st.form_submit_button("Filtrar Créditos")

            if filtrar_creditos_submit:
                df_credito = filtrar_creditos(st.session_state["df_extrato"])
                st.session_state["df_credito"] = df_credito
                st.markdown("#### Resultado (Extrato de Crédito)")
                st.dataframe(df_credito, use_container_width=True)

                pdf_creditos = df_to_pdf_bytes(df_credito, titulo="Extrato de Créditos")
                st.download_button(
                    label="Baixar PDF (Créditos)",
                    data=pdf_creditos,
                    file_name=f"creditos_{nome_cliente}.pdf",
                    mime="application/pdf",
                )

            if st.session_state.get("df_credito") is not None and not st.session_state["df_credito"].empty:
                with st.form("filtrar_glossario_credito_form"):
                    st.markdown("### 2) Filtrar Créditos no Glossário (com Precisão Ajustável)")
                    precision_credito = st.slider(
                        "Precisão da correspondência para Créditos (0.5 a 1.0):",
                        min_value=0.5,
                        max_value=1.0,
                        value=0.85,
                        step=0.025
                    )
                    filtrar_gloss_credito_submit = st.form_submit_button("Filtrar Créditos no Glossário")

                if filtrar_gloss_credito_submit:
                    df_credito_gloss = filtrar_por_glossario(
                        st.session_state["df_credito"], glossary_terms, threshold=int(precision_credito * 100)
                    )
                    df_credito_gloss = df_credito_gloss.drop(columns=["Débito (R$)", "Saldo (R$)"], errors='ignore')
                    st.session_state["df_credito_gloss"] = df_credito_gloss
                    st.session_state["df_credito_gloss_filtrado"] = None

                    st.markdown("#### Resultado: Créditos + Glossário")
                    st.dataframe(df_credito_gloss, use_container_width=True)

                    pdf_gloss_credito = df_to_pdf_bytes(df_credito_gloss, titulo="Créditos (Filtrados no Glossário)")
                    st.download_button(
                        label="Baixar PDF (Créditos Glossário)",
                        data=pdf_gloss_credito,
                        file_name=f"creditos_glossario_{nome_cliente}.pdf",
                        mime="application/pdf",
                    )

            if st.session_state.get("df_credito_gloss") is not None and not st.session_state["df_credito_gloss"].empty:
                with st.form("excluir_creditos_form"):
                    st.markdown("### 3) Lista Única de 'Histórico' para Créditos + Inclusão")

                    df_gloss_original_credito = st.session_state["df_credito_gloss"]
                    df_base_exclusao_credito = (
                        st.session_state["df_credito_gloss_filtrado"]
                        if st.session_state["df_credito_gloss_filtrado"] is not None
                        else df_gloss_original_credito
                    ).copy()

                    valores_unicos_credito = sorted(df_base_exclusao_credito["Histórico"].unique())
                    st.markdown("#### Lista Única de 'Histórico' (Créditos - sem repetições)")
                    st.write("Marque os itens que deseja incluir:")

                    selected_historicos_credito = []
                    for i, hist in enumerate(valores_unicos_credito):
                        count_hist = df_base_exclusao_credito[df_base_exclusao_credito["Histórico"] == hist].shape[0]
                        rotulo = f"{i+1}- {hist} ({count_hist} {'vez' if count_hist == 1 else 'vezes'})"
                        if st.checkbox(rotulo, key=f"unique_hist_credito_{i}"):
                            selected_historicos_credito.append(hist)

                    confirmar_inclusao_credito = st.form_submit_button("Confirmar Inclusão (Créditos)")
                    if confirmar_inclusao_credito:
                        if selected_historicos_credito:
                            df_filtrado_credito = df_base_exclusao_credito[
                                df_base_exclusao_credito["Histórico"].isin(selected_historicos_credito)
                            ]
                            df_filtrado_credito = df_filtrado_credito.reset_index(drop=True)

                            st.success("Operações de Crédito incluídas com sucesso!")
                            st.session_state["df_credito_gloss_filtrado"] = df_filtrado_credito

                            st.markdown("#### Lista Restante após Inclusões (Créditos - sem repetições)")
                            if df_filtrado_credito.empty:
                                st.write("Nenhum histórico de Crédito restante.")
                            else:
                                df_restante_unicos_credito = df_filtrado_credito["Histórico"].value_counts().reset_index()
                                df_restante_unicos_credito.columns = ["Histórico", "Ocorrências"]
                                st.dataframe(df_restante_unicos_credito, use_container_width=True)
                        else:
                            st.warning("Nenhuma descrição de Crédito foi selecionada.")

            if st.session_state.get("df_credito_gloss_filtrado") is not None and not st.session_state["df_credito_gloss_filtrado"].empty:
                with st.form("apresentar_tarifas_credito_form"):
                    st.markdown("### 4) Apresentar Tarifas para Créditos (DataFrame Final Ordenado)")
                    apresentar_tarifas_credito_submit = st.form_submit_button("Apresentar Tarifas para Créditos")

                if apresentar_tarifas_credito_submit:
                    df_para_exibir_credito = st.session_state["df_credito_gloss_filtrado"]
                    if not df_para_exibir_credito.empty:
                        df_para_exibir_credito = df_para_exibir_credito.drop(columns=["Débito (R$)", "Saldo (R$)"], errors='ignore')

                        # Converter para float (e forçar positivo)
                        numeros_credito = (
                            df_para_exibir_credito["Crédito (R$)"]
                            .str.replace("R$", "", regex=False)
                            .str.replace(" ", "", regex=False)
                            .str.replace(".", "", regex=False)
                            .str.replace(",", ".", regex=False)
                        )
                        valores_float_credito = pd.to_numeric(numeros_credito, errors='coerce').fillna(0.0)
                        valores_float_credito = valores_float_credito.abs()  # Sempre positivo

                        # Soma total
                        total_creditos = valores_float_credito.sum()

                        # Formatar em BRL
                        df_para_exibir_credito["Crédito (R$)"] = valores_float_credito.apply(formatar_valor_brl)

                        # Linhas de totais
                        valor_total_credito = pd.DataFrame({
                            "Data": [""],
                            "Histórico": ["Valor Total (R$)"],
                            "Docto.": [""],
                            "Crédito (R$)": [formatar_valor_brl(total_creditos)]
                        })
                        em_dobro_credito = pd.DataFrame({
                            "Data": [""],
                            "Histórico": ["Em dobro (R$)"],
                            "Docto.": [""],
                            "Crédito (R$)": [formatar_valor_brl(total_creditos * 2)]
                        })

                        extrato_credito_final = pd.concat([df_para_exibir_credito, valor_total_credito, em_dobro_credito], ignore_index=True)

                        # Ordenar cronologicamente
                        extrato_credito_final["Data"] = pd.to_datetime(
                            extrato_credito_final["Data"], format="%d/%m/%y", errors='coerce'
                        )
                        extrato_credito_final = extrato_credito_final.sort_values(by="Data")
                        extrato_credito_final["Data"] = extrato_credito_final["Data"].dt.strftime("%d/%m/%y")
                        extrato_credito_final["Data"] = extrato_credito_final["Data"].fillna("")

                        extrato_credito_final.loc[
                            extrato_credito_final["Histórico"].isin(["Valor Total (R$)", "Em dobro (R$)"]),
                            "Data"
                        ] = ""

                        st.markdown("#### DataFrame Final de Créditos (Cronológico)")
                        st.dataframe(extrato_credito_final, use_container_width=True)

                        pdf_final_credito = df_to_pdf_bytes(
                            extrato_credito_final,
                            titulo="Extrato Final de Créditos (Cronológico)",
                            formatar_linhas_especiais=True
                        )
                        st.download_button(
                            label="Baixar PDF (Créditos Final - Cronológico)",
                            data=pdf_final_credito,
                            file_name=f"creditos_final_cronologico_{nome_cliente}.pdf",
                            mime="application/pdf",
                        )

                        doc_final_credito = df_to_doc_bytes(
                            extrato_credito_final,
                            titulo="Extrato Final de Créditos (Cronológico)",
                            adicionar_totais=False
                        )
                        st.download_button(
                            label="Baixar DOCX (Créditos Final - Cronológico)",
                            data=doc_final_credito,
                            file_name=f"creditos_final_cronologico_{nome_cliente}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    else:
                        st.warning("Não há extrato final para Créditos.")

        # Botão para redefinir a operação
        if operacao in ["Débito", "Crédito"]:
            if st.button("Redefinir Seleção"):
                st.session_state["operacao_selecionada"] = None
                st.experimental_rerun()


if __name__ == "__main__":
    main()


